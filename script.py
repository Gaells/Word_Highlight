from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx2pdf import convert

def highlight_words_in_docx(file_path, words_to_highlight, output_path):
    doc = Document(file_path)

    def highlight_run(run, words_to_highlight):
        parts = []
        text = run.text
        pos = 0

        for word in words_to_highlight:
            lower_word = word.lower()
            start = text.lower().find(lower_word, pos)
            while start != -1:
                end = start + len(word)
                if start > pos:
                    parts.append((text[pos:start], run.font))
                parts.append((text[start:end], run.font, True))
                pos = end
                start = text.lower().find(lower_word, pos)
        
        if pos < len(text):
            parts.append((text[pos:], run.font))
        
        return parts

    def add_parts_to_paragraph(para, parts):
        for text, font, *highlight in parts:
            highlight = highlight[0] if highlight else False
            run = para.add_run(text)
            run.font.bold = font.bold
            run.font.italic = font.italic
            run.font.underline = font.underline
            run.font.size = font.size
            run.font.color.rgb = font.color.rgb
            run.font.name = font.name
            run.font.all_caps = font.all_caps
            run.font.complex_script = font.complex_script
            run.font.cs_bold = font.cs_bold
            run.font.cs_italic = font.cs_italic
            run.font.double_strike = font.double_strike
            run.font.emboss = font.emboss
            run.font.hidden = font.hidden
            run.font.highlight_color = font.highlight_color
            run.font.imprint = font.imprint
            run.font.math = font.math
            run.font.no_proof = font.no_proof
            run.font.outline = font.outline
            run.font.rtl = font.rtl
            run.font.shadow = font.shadow
            run.font.small_caps = font.small_caps
            run.font.snap_to_grid = font.snap_to_grid
            run.font.spec_vanish = font.spec_vanish
            run.font.strike = font.strike
            run.font.subscript = font.subscript
            run.font.superscript = font.superscript
            if highlight:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    def process_paragraphs(doc, words_to_highlight):
        for para in doc.paragraphs:
            original_runs = para.runs[:]
            para.clear()
            all_parts = []
            for run in original_runs:
                if run.element.xml.find('<w:checkBox') != -1:
                    # Adicionar o run do checkbox sem alterações
                    all_parts.append((run.text, run.font, False))
                else:
                    parts = highlight_run(run, words_to_highlight)
                    all_parts.extend(parts)
            add_parts_to_paragraph(para, all_parts)

    process_paragraphs(doc, words_to_highlight)
    doc.save(output_path)

def convert_docx_to_pdf(docx_path, pdf_path):
    convert(docx_path, pdf_path)

input_file = 'C:\\Users\\choqs\\Downloads\\Frm002 Supplier Self Evaluation Questionnaire Rev6 - Copiar.docx'
highlighted_file = 'modificado.docx'
pdf_file = 'highlight.pdf'
words = ['quality', 'division', 'plant:', 'street', 'address:', 'p.o.box no.:', 'city:', 'state:', 'zip:', 'telephone:', 'web site:', 'fax:', 'name:', 'title:', 'e-mail:']

highlight_words_in_docx(input_file, words, highlighted_file)
convert_docx_to_pdf(highlighted_file, pdf_file)
